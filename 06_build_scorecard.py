"""
06_build_scorecard.py — AFG Forecast Accuracy Index

Produces two elite outputs from resolved predictions in the scorecard database:
  1. reports/AFG_Scorecard_<date>.docx   — Word doc for records
  2. docs/index.html                      — GitHub Pages public dashboard

INTEGRITY CHECK: Prints DB state before building anything.
Every "Resolved" call must correspond to an actual Kalshi settlement.
Run 05_update_scorecard.py first to pull live resolution data.
"""

import os
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db import get_conn

# ── brand colours ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x17, 0x36, 0x5D)
BLUE        = RGBColor(0x36, 0x5F, 0x91)
GREEN       = RGBColor(0x1F, 0x6F, 0x3D)
RED         = RGBColor(0xA6, 0x23, 0x1F)
GRAY        = RGBColor(0x80, 0x80, 0x80)
HEADER_FILL = "DCE6F1"
GREEN_FILL  = "E8F5E9"
RED_FILL    = "FDECEA"
FONT        = "Cambria"

CATEGORY_ORDER = ["Sports", "Economics", "Politics", "Culture", "Weather"]
DOCX_DIR       = "reports"
WEB_DIR        = "docs"


# ── data ─────────────────────────────────────────────────────────────────────

def load_resolved():
    """
    Loads resolved predictions, applying the FIRST-CALL-ONLY rule:
    when the same market (same kalshi_ticker) was published across multiple
    report cycles, only AFG's earliest forecast counts toward the published
    Forecast Accuracy Index. This is the hardest test — it scores the call
    made at the point of maximum uncertainty, not one refined as evidence
    accumulated.

    Later republished calls on the same ticker remain in the database for
    internal reference but are excluded from all scorecard metrics.
    """
    with get_conn() as conn:
        rows = [dict(r) for r in conn.execute(
            "SELECT * FROM predictions WHERE status='Resolved' ORDER BY contract_close_date DESC"
        ).fetchall()]

    # Keep only the earliest report_date per ticker (fall back to market name
    # if a ticker is missing, so untickered legacy rows still dedupe sensibly).
    first_calls = {}
    duplicates_dropped = 0
    for r in rows:
        key = r.get("kalshi_ticker") or r.get("market")
        existing = first_calls.get(key)
        if existing is None:
            first_calls[key] = r
        else:
            duplicates_dropped += 1
            # keep whichever has the earlier report_date
            if str(r.get("report_date") or "") < str(existing.get("report_date") or ""):
                first_calls[key] = r

    deduped = sorted(
        first_calls.values(),
        key=lambda x: str(x.get("contract_close_date") or ""),
        reverse=True,
    )

    if duplicates_dropped:
        print(f"  Deduplication: {duplicates_dropped} repeat forecast(s) excluded "
              f"(first-call-only rule). {len(deduped)} unique market(s) scored.")

    return deduped


def call_correct(r):
    if r["recommendation"] == "BUY YES":
        return r["outcome"] == 1
    if r["recommendation"] == "BUY NO":
        return r["outcome"] == 0
    return None


def compute_metrics(resolved):
    n = len(resolved)
    if n == 0:
        return None

    correct_list = [r for r in resolved if call_correct(r) is True]
    incorrect_list = [r for r in resolved if call_correct(r) is False]

    afg_brier    = sum(r["brier_score"] for r in resolved) / n
    kalshi_brier = sum(r["kalshi_brier_score"] for r in resolved) / n

    by_cat = {}
    for cat in CATEGORY_ORDER:
        subset = [r for r in resolved if r["category"] == cat]
        if not subset:
            continue
        c = [r for r in subset if call_correct(r) is True]
        ic = [r for r in subset if call_correct(r) is False]
        by_cat[cat] = {
            "n":        len(subset),
            "correct":  len(c),
            "incorrect": len(ic),
            "accuracy": len(c) / len(subset) if subset else None,
            "afg_brier": sum(r["brier_score"] for r in subset) / len(subset),
        }

    return {
        "n":            n,
        "correct":      len(correct_list),
        "incorrect":    len(incorrect_list),
        "accuracy":     len(correct_list) / n,
        "afg_brier":    afg_brier,
        "kalshi_brier": kalshi_brier,
        "skill_delta":  kalshi_brier - afg_brier,
        "by_cat":       by_cat,
    }


# ── docx helpers ─────────────────────────────────────────────────────────────

def _run(p, text, size=10.5, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(size)
    r.font.bold = bold; r.font.italic = italic
    if color: r.font.color.rgb = color
    return r


def _shade(cell, fill):
    s = OxmlElement("w:shd")
    s.set(qn("w:fill"), fill); s.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(s)


def _borders(cell, color="B0B8C8"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for e in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "2")
        el.set(qn("w:color"), color)
        b.append(el)
    tcPr.append(b)


def _set_widths(t, widths):
    tblPr = t._tbl.tblPr
    for x in tblPr.findall(qn("w:tblLayout")): tblPr.remove(x)
    for x in tblPr.findall(qn("w:tblW")):     tblPr.remove(x)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(sum(widths))); w.set(qn("w:type"), "dxa")
    tblPr.append(w)
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is not None: t._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for wi in widths:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(wi)); grid.append(gc)
    t._tbl.insert(list(t._tbl).index(tblPr) + 1, grid)
    for row in t.rows:
        for i, c in enumerate(row.cells):
            if i < len(widths):
                tcPr = c._tc.get_or_add_tcPr()
                for old in tcPr.findall(qn("w:tcW")): tcPr.remove(old)
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def _heading(doc, text, size=13, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    _run(p, text, size=size, bold=True, color=color)
    pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), "365F91"); bot.set(qn("w:space"), "3")
    pbdr.append(bot); p._p.get_or_add_pPr().append(pbdr)


def _metric_row(doc, label, value, value_color=None, sub=None):
    """One labelled metric line — big value, small label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(1)
    _run(p, f"{label}: ", size=10, bold=True, color=NAVY)
    _run(p, str(value), size=14, bold=True, color=value_color or NAVY)
    if sub:
        _run(p, f"  {sub}", size=9, italic=True, color=GRAY)


# ── Word doc ─────────────────────────────────────────────────────────────────

def build_docx(resolved, m, out_path):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Inches(8.5); sec.page_height = Inches(11)
    for edge in ("left","right","top","bottom"):
        setattr(sec, f"{edge}_margin", Inches(0.75))
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10.5)

    # Title block
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    _run(p, "Axiom Forecasting Group", size=26, bold=True, color=NAVY)
    pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"14")
    bot.set(qn("w:color"),"17365D"); bot.set(qn("w:space"),"4")
    pbdr.append(bot); p._p.get_or_add_pPr().append(pbdr)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(5); p2.paragraph_format.space_after = Pt(14)
    _run(p2, "AFG Forecast Accuracy Index", size=14, bold=True, color=NAVY)
    _run(p2, f"   |   Updated {date.today().strftime('%B %d, %Y')}", size=12, color=BLUE)

    if m is None:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_before = Pt(24)
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p3, "No resolved predictions yet.", size=12, bold=True, color=GRAY)
        p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p4, "The Forecast Accuracy Index populates as markets settle on Kalshi.",
             size=10.5, italic=True, color=GRAY)
        doc.save(out_path)
        return

    # ── Overall metrics ───────────────────────────────────────────────────────
    _heading(doc, "Overall Performance")

    beat = m["skill_delta"] > 0
    # 4-cell metric grid
    t = doc.add_table(rows=1, cols=4)
    labels   = ["Total Forecasts", "Correct", "Incorrect", "AFG Brier Score"]
    values   = [m["n"], m["correct"], m["incorrect"], f"{m['afg_brier']:.3f}"]
    colors   = [NAVY, GREEN, RED, NAVY]
    fills    = ["F8FAFC", GREEN_FILL, RED_FILL, "F8FAFC"]
    for i, (lbl, val, col, fill) in enumerate(zip(labels, values, colors, fills)):
        c = t.rows[0].cells[i]; c.text = ""
        _shade(c, fill); _borders(c, "C8D4E8")
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(8); cp.paragraph_format.space_after = Pt(2)
        _run(cp, str(val), size=22, bold=True, color=col)
        p_lbl = c.add_paragraph(); p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_lbl.paragraph_format.space_before = Pt(0); p_lbl.paragraph_format.space_after = Pt(8)
        _run(p_lbl, lbl, size=9, bold=True, color=GRAY)
    _set_widths(t, [2700, 2700, 2700, 2700])

    # Skill delta line
    p_sk = doc.add_paragraph()
    p_sk.paragraph_format.space_before = Pt(8); p_sk.paragraph_format.space_after = Pt(2)
    p_sk.alignment = WD_ALIGN_PARAGRAPH.CENTER
    delta_word = "beats" if beat else "trails"
    delta_col  = GREEN if beat else RED
    _run(p_sk, "Accuracy: ", size=11, bold=True, color=NAVY)
    _run(p_sk, f"{m['accuracy']:.0%}", size=13, bold=True, color=NAVY)
    _run(p_sk, "   |   AFG ", size=10.5, color=GRAY)
    _run(p_sk, f"{delta_word} the Kalshi market", size=10.5, bold=True, color=delta_col)
    _run(p_sk, f" by {abs(m['skill_delta']):.3f} Brier points", size=10.5, color=GRAY)
    _run(p_sk, f"  (Market Brier: {m['kalshi_brier']:.3f})", size=9.5, italic=True, color=GRAY)

    # ── By category ───────────────────────────────────────────────────────────
    _heading(doc, "Accuracy by Category")
    cat_t = doc.add_table(rows=1, cols=6)
    headers = ["Category", "Forecasts", "Correct", "Incorrect", "Accuracy", "Brier Score"]
    for i, h in enumerate(headers):
        c = cat_t.rows[0].cells[i]; c.text = ""
        cp = c.paragraphs[0]
        cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
        _run(cp, h, size=9.5, bold=True, color=NAVY)
        _shade(c, HEADER_FILL); _borders(c)

    for cat in CATEGORY_ORDER:
        d = m["by_cat"].get(cat)
        row = cat_t.add_row()
        vals = [
            cat,
            str(d["n"])         if d else "—",
            str(d["correct"])   if d else "—",
            str(d["incorrect"]) if d else "—",
            f"{d['accuracy']:.0%}" if d and d["accuracy"] is not None else "—",
            f"{d['afg_brier']:.3f}" if d else "—",
        ]
        colors_row = [NAVY, GRAY, GREEN, RED, NAVY, NAVY]
        for i, (val, col) in enumerate(zip(vals, colors_row)):
            c = row.cells[i]; c.text = ""
            cp = c.paragraphs[0]
            cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
            _run(cp, val, size=9.5, bold=(i == 0), color=col if val != "—" else GRAY)
            _borders(c)
    _set_widths(cat_t, [2100, 1500, 1500, 1500, 1500, 1500])

    # ── Recent resolved calls ────────────────────────────────────────────────
    if resolved:
        _heading(doc, "Recent Resolved Forecasts")
        rec_t = doc.add_table(rows=1, cols=6)
        r_headers = ["Market", "Category", "Kalshi", "AFG", "Outcome", "Result"]
        for i, h in enumerate(r_headers):
            c = rec_t.rows[0].cells[i]; c.text = ""
            cp = c.paragraphs[0]
            cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
            _run(cp, h, size=9.5, bold=True, color=NAVY)
            _shade(c, HEADER_FILL); _borders(c)

        for r in resolved[:15]:
            hit = call_correct(r)
            row = rec_t.add_row()
            result_text  = "✓ Correct"   if hit else "✗ Incorrect"
            result_color = GREEN         if hit else RED
            outcome_txt  = "YES"         if r["outcome"] == 1 else "NO"
            vals = [
                r["market"][:48],
                r["category"],
                f"{r['kalshi_price']:.0%}",
                f"{r['afg_probability']:.0%}",
                outcome_txt,
                result_text,
            ]
            cell_colors = [NAVY, GRAY, GRAY, NAVY, GRAY, result_color]
            for i, (val, col) in enumerate(zip(vals, cell_colors)):
                c = row.cells[i]; c.text = ""
                cp = c.paragraphs[0]
                cp.paragraph_format.space_before = Pt(3); cp.paragraph_format.space_after = Pt(3)
                _run(cp, val, size=9, bold=(i == 5), color=col)
                _borders(c)
        _set_widths(rec_t, [4400, 1300, 900, 900, 900, 1400])

    # footnote
    p_fn = doc.add_paragraph()
    p_fn.paragraph_format.space_before = Pt(16)
    _run(p_fn, "Brier Score: 0 = perfect · 0.25 = coin flip · 1.0 = perfectly wrong.  "
               "This report is for informational purposes only and is not financial advice.",
         size=8.5, italic=True, color=GRAY)

    doc.save(out_path)


# ── Website ───────────────────────────────────────────────────────────────────

def build_website(resolved, m, out_dir):
    os.makedirs(out_dir, exist_ok=True)
    updated = date.today().strftime("%B %d, %Y")
    year    = date.today().year

    if m is None:
        body = """
      <div class="empty-state">
        <div class="empty-icon">📊</div>
        <h2>No resolved forecasts yet</h2>
        <p>The AFG Forecast Accuracy Index populates automatically as markets settle on Kalshi.<br>
           Check back after the first resolution date passes.</p>
      </div>"""
    else:
        beat = m["skill_delta"] > 0
        delta_cls  = "good" if beat else "bad"
        delta_word = "outperforms" if beat else "trails"

        # headline cards
        cards = f"""
      <div class="cards">
        <div class="card">
          <div class="card-value">{m['n']}</div>
          <div class="card-label">Total Forecasts Issued</div>
        </div>
        <div class="card good-card">
          <div class="card-value good">{m['correct']}</div>
          <div class="card-label">Correct Forecasts</div>
        </div>
        <div class="card bad-card">
          <div class="card-value bad">{m['incorrect']}</div>
          <div class="card-label">Incorrect Forecasts</div>
        </div>
        <div class="card">
          <div class="card-value">{m['accuracy']:.0%}</div>
          <div class="card-label">Overall Accuracy</div>
        </div>
        <div class="card">
          <div class="card-value">{m['afg_brier']:.3f}</div>
          <div class="card-label">AFG Brier Score</div>
          <div class="card-sub">lower is better</div>
        </div>
        <div class="card muted-card">
          <div class="card-value muted">{m['kalshi_brier']:.3f}</div>
          <div class="card-label">Market Brier Score</div>
          <div class="card-sub">the crowd's accuracy</div>
        </div>
      </div>
      <p class="skill-line {delta_cls}">
        AFG {delta_word} the Kalshi market by
        <strong>{abs(m['skill_delta']):.3f} Brier points</strong>
        across {m['n']} resolved forecasts.
      </p>"""

        # category table
        cat_rows = ""
        for cat in CATEGORY_ORDER:
            d = m["by_cat"].get(cat)
            if d:
                acc_pct = f"{d['accuracy']:.0%}" if d["accuracy"] is not None else "—"
                beat_cat = d["afg_brier"] < 0.25
                brier_cls = "good" if beat_cat else ""
                cat_rows += f"""
          <tr>
            <td class="cat-name">{cat}</td>
            <td>{d['n']}</td>
            <td class="good">{d['correct']}</td>
            <td class="bad">{d['incorrect']}</td>
            <td><strong>{acc_pct}</strong></td>
            <td class="{brier_cls}">{d['afg_brier']:.3f}</td>
          </tr>"""
            else:
                cat_rows += f"<tr><td class='cat-name'>{cat}</td><td colspan='5' class='muted'>No resolved forecasts yet</td></tr>"

        # recent resolved
        recent_rows = ""
        for r in resolved[:20]:
            hit = call_correct(r)
            hit_cls  = "good" if hit else "bad"
            hit_icon = "✓" if hit else "✗"
            outcome  = "YES" if r["outcome"] == 1 else "NO"
            recent_rows += f"""
          <tr>
            <td>{r['market'][:60]}</td>
            <td class="cat-badge">{r['category']}</td>
            <td>{r['kalshi_price']:.0%}</td>
            <td><strong>{r['afg_probability']:.0%}</strong></td>
            <td>{outcome}</td>
            <td class="{hit_cls} result-cell">{hit_icon} {"Correct" if hit else "Incorrect"}</td>
          </tr>"""

        body = f"""
      {cards}

      <section>
        <h2>Accuracy by Category</h2>
        <table>
          <thead>
            <tr>
              <th>Category</th><th>Forecasts</th><th>Correct</th>
              <th>Incorrect</th><th>Accuracy</th><th>Brier Score</th>
            </tr>
          </thead>
          <tbody>{cat_rows}</tbody>
        </table>
      </section>

      <section>
        <h2>Recent Resolved Forecasts</h2>
        <div class="table-wrap">
        <table>
          <thead>
            <tr>
              <th>Market</th><th>Category</th><th>Kalshi</th>
              <th>AFG</th><th>Outcome</th><th>Result</th>
            </tr>
          </thead>
          <tbody>{recent_rows}</tbody>
        </table>
        </div>
      </section>"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>AFG Forecast Accuracy Index</title>
  <style>
    :root {{
      --navy:  #17365D;
      --blue:  #365F91;
      --green: #1F6F3D;
      --red:   #A6231F;
      --green-bg: #E8F5E9;
      --red-bg:   #FDECEA;
      --gray:  #6B7280;
      --line:  #E5EAF0;
      --bg:    #F8FAFC;
      --white: #FFFFFF;
    }}
    *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{
      font-family: Georgia, 'Cambria', serif;
      background: var(--bg);
      color: #1A202C;
      line-height: 1.6;
    }}

    /* ── Header ── */
    header {{
      background: var(--navy);
      color: white;
      padding: 36px 24px 28px;
      text-align: center;
    }}
    .logo-rule {{
      width: 60px; height: 3px;
      background: var(--blue);
      margin: 0 auto 16px;
      border-radius: 2px;
    }}
    header h1 {{ font-size: 2rem; letter-spacing: 0.02em; font-weight: 700; }}
    header .sub {{ font-size: 1rem; color: #93B4D4; margin-top: 6px; letter-spacing: 0.05em; text-transform: uppercase; }}
    header .updated {{ font-size: 0.82rem; color: #6A8FAF; margin-top: 8px; }}

    /* ── Main ── */
    main {{ max-width: 1040px; margin: 0 auto; padding: 36px 20px 60px; }}

    /* ── Metric cards ── */
    .cards {{
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(148px, 1fr));
      gap: 14px;
      margin-bottom: 20px;
    }}
    .card {{
      background: var(--white);
      border: 1px solid var(--line);
      border-radius: 12px;
      padding: 20px 16px 16px;
      text-align: center;
      box-shadow: 0 1px 4px rgba(0,0,0,.05);
    }}
    .good-card {{ border-color: #A8D5B5; background: var(--green-bg); }}
    .bad-card  {{ border-color: #F0B8B5; background: var(--red-bg); }}
    .muted-card {{ background: #F1F5F9; }}
    .card-value  {{ font-size: 2.2rem; font-weight: 700; color: var(--navy); line-height: 1.1; }}
    .card-label  {{ font-size: 0.78rem; color: var(--gray); text-transform: uppercase; letter-spacing: 0.06em; margin-top: 6px; }}
    .card-sub    {{ font-size: 0.72rem; color: #9CA3AF; margin-top: 2px; }}
    .good  {{ color: var(--green) !important; }}
    .bad   {{ color: var(--red)   !important; }}
    .muted {{ color: var(--gray)  !important; }}

    /* ── Skill line ── */
    .skill-line {{
      font-size: 1.05rem;
      text-align: center;
      margin: 4px 0 32px;
      padding: 12px 20px;
      border-radius: 8px;
      background: var(--white);
      border: 1px solid var(--line);
    }}
    .skill-line.good {{ color: var(--green); border-color: #A8D5B5; background: var(--green-bg); }}
    .skill-line.bad  {{ color: var(--red);   border-color: #F0B8B5; background: var(--red-bg); }}

    /* ── Section headings ── */
    section {{ margin-top: 36px; }}
    h2 {{
      font-size: 1.1rem;
      font-weight: 700;
      color: var(--navy);
      border-bottom: 2px solid var(--blue);
      padding-bottom: 6px;
      margin-bottom: 14px;
      letter-spacing: 0.01em;
    }}

    /* ── Tables ── */
    .table-wrap {{ overflow-x: auto; }}
    table {{
      border-collapse: collapse;
      width: 100%;
      font-size: 0.92rem;
      background: var(--white);
      border-radius: 8px;
      overflow: hidden;
      box-shadow: 0 1px 4px rgba(0,0,0,.04);
    }}
    th {{
      background: #DCE6F1;
      color: var(--navy);
      font-weight: 700;
      text-align: left;
      padding: 10px 14px;
      border-bottom: 2px solid var(--blue);
      font-size: 0.85rem;
      text-transform: uppercase;
      letter-spacing: 0.04em;
    }}
    td {{
      padding: 10px 14px;
      border-bottom: 1px solid var(--line);
      vertical-align: middle;
    }}
    tbody tr:last-child td {{ border-bottom: none; }}
    tbody tr:hover {{ background: #F0F4FA; }}
    .cat-name  {{ font-weight: 700; color: var(--navy); }}
    .cat-badge {{
      font-size: 0.78rem;
      background: #EEF2F8;
      color: var(--blue);
      padding: 2px 8px;
      border-radius: 12px;
      white-space: nowrap;
      font-weight: 600;
    }}
    .result-cell {{ font-weight: 700; font-size: 0.9rem; }}

    /* ── Empty state ── */
    .empty-state {{
      text-align: center;
      padding: 64px 20px;
      color: var(--gray);
    }}
    .empty-icon {{ font-size: 3rem; margin-bottom: 16px; }}
    .empty-state h2 {{ font-size: 1.3rem; color: var(--navy); margin-bottom: 10px; }}
    .empty-state p  {{ font-size: 0.95rem; line-height: 1.7; }}

    /* ── Footer ── */
    footer {{
      background: var(--navy);
      color: #6A8FAF;
      font-size: 0.82rem;
      text-align: center;
      padding: 24px 20px;
      margin-top: 60px;
      line-height: 1.7;
    }}
    footer strong {{ color: #93B4D4; }}

    /* ── Mobile ── */
    @media (max-width: 640px) {{
      header h1 {{ font-size: 1.4rem; }}
      .cards {{ grid-template-columns: repeat(2, 1fr); }}
      .card-value {{ font-size: 1.7rem; }}
    }}
  </style>
</head>
<body>
  <header>
    <div class="logo-rule"></div>
    <h1>AFG Forecast Accuracy Index</h1>
    <div class="sub">Axiom Forecasting Group · Prediction Markets Research</div>
    <div class="updated">Last updated {updated}</div>
  </header>

  <main>
    {body}
  </main>

  <footer>
    <strong>Brier Score:</strong> 0 = perfect &nbsp;·&nbsp; 0.25 = coin flip &nbsp;·&nbsp;
    1.0 = perfectly wrong &nbsp;·&nbsp; Lower is better.<br>
    AFG publishes calibrated probability estimates on Kalshi prediction markets every Monday, Wednesday, and Friday.<br>
    This index is for informational purposes only and does not constitute financial advice.
    &copy; {year} Axiom Forecasting Group
  </footer>
</body>
</html>"""

    with open(os.path.join(out_dir, "index.html"), "w", encoding="utf-8") as f:
        f.write(html)


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(DOCX_DIR, exist_ok=True)

    # INTEGRITY CHECK
    with get_conn() as conn:
        n_open     = conn.execute("SELECT COUNT(*) FROM predictions WHERE status='Open'").fetchone()[0]
        n_resolved = conn.execute("SELECT COUNT(*) FROM predictions WHERE status='Resolved'").fetchone()[0]
    print(f"=== SCORECARD INTEGRITY CHECK ===")
    print(f"  Total logged    : {n_open + n_resolved}")
    print(f"  Open            : {n_open}")
    print(f"  Resolved        : {n_resolved}")
    if n_resolved > 0:
        print(f"  WARNING: {n_resolved} resolved calls will be published.")
        print(f"  Confirm each against actual Kalshi settlement before releasing.")
    else:
        print(f"  Output will show: no resolved forecasts yet.")
    print(f"=================================")

    resolved = load_resolved()
    m        = compute_metrics(resolved)

    docx_path = f"{DOCX_DIR}/AFG_Scorecard_{date.today().isoformat()}.docx"
    build_docx(resolved, m, docx_path)
    build_website(resolved, m, WEB_DIR)

    print(f"Wrote {docx_path}")
    print(f"Wrote {WEB_DIR}/index.html")
    if m:
        print(f"  {m['n']} resolved | AFG Brier {m['afg_brier']:.3f} vs market {m['kalshi_brier']:.3f} | Accuracy {m['accuracy']:.0%}")
    else:
        print(f"  No resolved forecasts — index will populate as markets settle.")


if __name__ == "__main__":
    main()
