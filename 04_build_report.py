"""
04_build_report.py — Phase 4. Reads the two files YOU approved after review
(data/approved_predictions.csv and data/approved_narrative.md) and produces:

  - reports/AFG_Research_Report_<date>.docx  (locked institutional template)
  - data/final_report.xlsx                    (raw data backup)
  - reports/AFG_Intelligence_Brief_<date>_Substack.md  (public-facing brief)

This script refuses to run if either approved file is missing -- that's the
checkpoint gate. Nothing before this point is allowed to skip your review.

Run:
    python 04_build_report.py
"""

import os
import re
import sys
from datetime import date

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PREDICTIONS_PATH = "data/approved_predictions.csv"
NARRATIVE_PATH = "data/approved_narrative.md"
XLSX_OUTPUT_PATH = "data/final_report.xlsx"
REPORTS_DIR = "reports"

NAVY = RGBColor(0x17, 0x36, 0x5D)
BLUE = RGBColor(0x36, 0x5F, 0x91)
ACCENT = RGBColor(0x4F, 0x81, 0xBD)
GREEN = RGBColor(0x1F, 0x6F, 0x3D)
RED = RGBColor(0xA6, 0x23, 0x1F)
GRAY = RGBColor(0x80, 0x80, 0x80)
HEADER_FILL = "DCE6F1"
BORDER_COLOR = "A6A6A6"

FONT = "Cambria"
CATEGORY_ORDER = ["Sports", "Politics", "Economics", "Culture", "Weather"]

# Explicit DXA column widths (1440 dxa = 1 inch). Landscape Letter with 0.5"
# margins gives ~13680 dxa of usable width. The conviction table is indented
# -2880 dxa (2 inches left of the body margin) per the locked format.
# Column widths constrained to fit within the 10800 DXA text column
# (page 12240, margins 720 each side). The 7/17 reference used 14400 DXA
# which overflows in Word but clips in LibreOffice. These fit correctly in both.
# DXA widths fitted to 10800 DXA text column (portrait Letter, 0.5" margins).
# Recommendation: 2000 DXA -- enough for the full word "Recommendation" in 9.5pt bold Cambria.
# Market: 3200 DXA -- narrow enough to let the row fit on one page.
# Kalshi/AFG/Edge: 800 DXA each -- short numeric content.
OPP_COL_WIDTHS    = [3200, 1500, 800, 800, 800, 1700, 2000]      # sum=10800
CONTRA_COL_WIDTHS = [1440, 3060, 3780, 2520]                      # sum=10800
DASH_COL_WIDTHS   = [3600, 7200]                                   # sum=10800
TABLE_LEFT_SHIFT  = 0   # tables sized to fit text column; no indent needed



# ---------- input loading & validation ----------

def require_approved_files():
    missing = [p for p in (PREDICTIONS_PATH, NARRATIVE_PATH) if not os.path.exists(p)]
    if missing:
        print("ERROR: Phase 4 requires both approved files to exist. Missing:")
        for m in missing:
            print(f"  {m}")
        print("\nThis is the review gate -- run 03_run_claude.py, fill in the draft")
        print("files, review them, and save your approved versions with these exact")
        print("filenames before running Phase 4.")
        sys.exit(1)


def load_predictions():
    df = pd.read_csv(PREDICTIONS_PATH)
    required = ["market", "category", "kalshi_price", "afg_probability", "conviction", "recommendation"]
    missing_cols = [c for c in required if c not in df.columns]
    if missing_cols:
        raise ValueError(f"{PREDICTIONS_PATH} is missing required columns: {missing_cols}")

    df["edge_score"] = df["afg_probability"] - df["kalshi_price"]
    df = df.sort_values("edge_score", key=lambda s: s.abs(), ascending=False).reset_index(drop=True)
    return df


def parse_narrative():
    """
    Splits approved_narrative.md on '## ' headers into a dict of section name
    -> raw text. Also parses the Contrarian Positioning markdown table into a
    list of row dicts.
    """
    with open(NARRATIVE_PATH, encoding="utf-8") as f:
        text = f.read()

    sections = {}
    parts = re.split(r"^## (.+)$", text, flags=re.MULTILINE)
    # re.split with a capturing group interleaves: [prefix, header1, body1, header2, body2, ...]
    for i in range(1, len(parts), 2):
        header = parts[i].strip()
        body = parts[i + 1].strip() if i + 1 < len(parts) else ""
        sections[header] = body

    contrarian_rows = []
    if "Contrarian Positioning" in sections:
        table_text = sections["Contrarian Positioning"]
        lines = [l.strip() for l in table_text.splitlines() if l.strip().startswith("|")]
        for line in lines[2:]:  # skip header row and separator row
            cells = [c.strip() for c in line.strip("|").split("|")]
            if len(cells) >= 4:
                contrarian_rows.append({
                    "category": cells[0], "consensus": cells[1],
                    "afg_view": cells[2], "position": cells[3],
                })

    return sections, contrarian_rows


# ---------- docx helpers ----------

def set_cell_shading(cell, hex_color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border_all(cell, color=BORDER_COLOR, size=2):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_table_layout(table, col_widths_dxa, indent_dxa=0):
    """
    Matched byte-for-byte to the 2026-07-17 reference report:
      - NO table.autofit = False  (that adds tblLayout=fixed which clips cells)
      - NO tblInd
      - NO tblLayout element at all
      - Sets tblW to total DXA
      - Rebuilds tblGrid with per-column widths
      - Sets tcW on every cell

    The 7/17 reference has none of the "fixed layout" machinery that causes
    columns to clip their content -- it relies purely on tblW + gridCol + tcW.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tblPr = table._tbl.tblPr

    # Remove tblLayout (added by python-docx when autofit=False; absent in 7/17)
    for old_el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(old_el)

    # Remove tblInd (absent in 7/17)
    for old_el in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old_el)

    # Set tblW
    total = sum(col_widths_dxa)
    for old_el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old_el)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Rebuild tblGrid
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_dxa:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    tblPr_idx = list(table._tbl).index(tblPr)
    table._tbl.insert(tblPr_idx + 1, grid)

    # Set minimal cell margins (top/bottom 40 DXA = 28pt/20) to compact row height
    tblCellMar = OxmlElement("w:tblCellMar")
    for edge, val in (("top","28"), ("bottom","28"), ("left","80"), ("right","80")):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")
        tblCellMar.append(el)
    tblPr.append(tblCellMar)

    # Set tcW on every cell
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx >= len(col_widths_dxa):
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            for old_el in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old_el)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(col_widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def style_run(run, size=10.5, bold=False, italic=False, color=None, underline=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color is not None:
        run.font.color.rgb = color
    return run


def add_title_block(doc, report_date):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Axiom Forecasting Group")
    style_run(run, size=22, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(2)

    # bottom rule under the title
    p_border = OxmlElement("w:pPr")
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "17365D")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p2.add_run("AFG Research Report")
    style_run(r1, size=12, bold=True, color=NAVY)
    r2 = p2.add_run(f"   |   {report_date}   |   Prediction Markets Research")
    style_run(r2, size=11, color=BLUE)
    p2.paragraph_format.space_after = Pt(8)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=14, bold=True, color=BLUE, underline=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body_paragraph(doc, text, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_run(run, size=size, italic=italic)
    p.paragraph_format.space_after = Pt(4)
    return p


def recommendation_color(rec):
    rec = (rec or "").upper()
    if "BUY YES" in rec:
        return GREEN
    if "BUY NO" in rec:
        return RED
    return GRAY


def conviction_color(conv):
    conv = (conv or "").upper()
    if conv == "HIGH":
        return NAVY
    if conv == "MEDIUM":
        return BLUE
    return GRAY


def add_opportunities_table(doc, df):
    headers = ["Market", "Category", "Kalshi", "AFG", "Edge", "Conviction", "Recommendation"]
    table = doc.add_table(rows=1, cols=len(headers))

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        style_run(run, size=9.5, bold=True, color=NAVY)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border_all(cell)

    for _, row in df.iterrows():
        tr = table.add_row()
        # Prevent any single row from splitting across a page break
        trPr = tr._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        cantSplit.set(qn("w:val"), "1")
        trPr.append(cantSplit)
        cells = tr.cells
        values = [
            row["market"],
            row["category"],
            f"{row['kalshi_price']:.0%}",
            f"{row['afg_probability']:.0%}",
            f"{row['edge_score']*100:+.0f}pp",
            row["conviction"],
            row["recommendation"],
        ]
        colors = [None, None, None, None, None, conviction_color(row["conviction"]), recommendation_color(row["recommendation"])]
        for i, (val, color) in enumerate(zip(values, colors)):
            cell = cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            style_run(run, size=9, bold=(color is not None), color=color)
            set_cell_border_all(cell)

    set_table_layout(table, OPP_COL_WIDTHS, indent_dxa=TABLE_LEFT_SHIFT)
    return table


def add_contrarian_table(doc, rows):
    headers = ["Category", "Consensus", "AFG View", "Position"]
    table = doc.add_table(rows=1, cols=len(headers))

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        style_run(run, size=9.5, bold=True, color=NAVY)
        set_cell_shading(cell, HEADER_FILL)
        set_cell_border_all(cell)

    for row in rows:
        cells = table.add_row().cells
        values = [row["category"], row["consensus"], row["afg_view"], row["position"]]
        for i, val in enumerate(values):
            cell = cells[i]
            cell.text = ""
            color = recommendation_color(val) if i == 3 else None
            run = cell.paragraphs[0].add_run(val)
            style_run(run, size=9.5, bold=(i == 0 or color is not None), color=color)
            set_cell_border_all(cell)

    set_table_layout(table, CONTRA_COL_WIDTHS, indent_dxa=TABLE_LEFT_SHIFT)
    return table


def set_landscape(doc):
    """
    Portrait US Letter with 0.5" margins, matching the 2026-07-17 reference
    report. (Name kept for compatibility.) Tables are 14400 DXA and pulled
    left via negative indent so they sit centered on the page, extending
    past the text column on both sides -- this is what the reference does.
    """
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)


# ---------- assembly ----------

def add_dashboard_table(doc, stats):
    table = doc.add_table(rows=0, cols=2)
    for label, value in stats:
        cells = table.add_row().cells
        c0 = cells[0]
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(label)
        style_run(r0, size=9.5, bold=True, color=NAVY)
        set_cell_border_all(c0)
        c1 = cells[1]
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(str(value))
        style_run(r1, size=9.5)
        set_cell_border_all(c1)
    set_table_layout(table, DASH_COL_WIDTHS, indent_dxa=TABLE_LEFT_SHIFT)
    return table


def build_docx(df, sections, contrarian_rows, report_date, output_path):
    doc = Document()
    set_landscape(doc)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)

    add_title_block(doc, report_date)

    add_section_heading(doc, "Executive Summary")
    add_body_paragraph(doc, sections.get("Executive Summary", ""))

    # Executive Dashboard stats, computed from the approved predictions
    n_published = len(df)
    n_actionable = int((df["edge_score"].abs() >= 0.03).sum())
    n_high = int((df["conviction"].str.upper() == "HIGH").sum())
    excluded_note = sections.get("Publication Note", "")
    # user records the screened/excluded counts in the narrative; fall back to published count
    add_section_heading(doc, "Executive Dashboard")
    add_dashboard_table(doc, [
        ("Markets screened", sections.get("Markets Screened", "25 (5 categories × 5)")),
        ("Markets excluded for data quality", 25 - n_published),
        ("Markets published", n_published),
        ("Markets with actionable edge (≥ 3pp)", n_actionable),
        ("HIGH conviction calls", n_high),
    ])

    add_section_heading(doc, "10 Highest Conviction Opportunities")
    add_body_paragraph(doc, "Top 10, ranked by absolute AFG Edge Score (AFG Probability minus Kalshi implied probability).", italic=True, size=9)
    add_opportunities_table(doc, df.head(10))

    for category in CATEGORY_ORDER:
        if category in sections and sections[category]:
            add_section_heading(doc, category)
            add_body_paragraph(doc, sections[category])

    if "Cross-Asset Themes" in sections:
        add_section_heading(doc, "Cross-Asset Themes")
        for line in sections["Cross-Asset Themes"].splitlines():
            line = line.strip().lstrip("-").strip()
            if line:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(line)
                style_run(run, size=10.5)

    if contrarian_rows:
        add_section_heading(doc, "Contrarian Positioning")
        add_contrarian_table(doc, contrarian_rows)

    if excluded_note:
        add_section_heading(doc, "Publication Note")
        add_body_paragraph(doc, excluded_note, italic=True, size=9.5)

    doc.save(output_path)


def build_xlsx(df, output_path, raw_markets_path="data/raw_markets.xlsx"):
    """
    Three sheets:
      - 'Analysis'             : the approved calls
      - 'Volume Verification'  : top 5 by TRADING VOLUME per category, taken
                                 straight from the Phase 1 pull, so you can
                                 open Kalshi's site side by side and confirm
                                 the rankings match.
      - 'Category Totals'      : total + top-5 volume per category.

    If raw_markets.xlsx is missing this now RAISES rather than silently
    skipping the volume sheets (which previously produced a workbook with no
    volume data and no warning).
    """
    export_df = df.copy()
    export_df["kalshi_price"] = export_df["kalshi_price"].map(lambda x: f"{x:.1%}")
    export_df["afg_probability"] = export_df["afg_probability"].map(lambda x: f"{x:.1%}")
    export_df["edge_score"] = export_df["edge_score"].map(lambda x: f"{x*100:+.1f}pp")

    if not os.path.exists(raw_markets_path):
        raise FileNotFoundError(
            f"{raw_markets_path} not found -- the Volume Verification sheet cannot be "
            f"built without it. Run 01_pull_markets.py before 04_build_report.py."
        )

    raw = pd.read_excel(raw_markets_path)
    if "volume_fp" not in raw.columns:
        raise ValueError(
            f"{raw_markets_path} has no 'volume_fp' column. Re-run 01_pull_markets.py."
        )

    raw["volume_fp"] = pd.to_numeric(raw["volume_fp"], errors="coerce").fillna(0)

    top5 = (
        raw.sort_values("volume_fp", ascending=False)
        .groupby("category", group_keys=False)
        .head(5)
        .sort_values(["category", "volume_fp"], ascending=[True, False])
        .reset_index(drop=True)
    )
    verify = pd.DataFrame({
        "Category": top5["category"],
        "Rank in Category": top5.groupby("category").cumcount() + 1,
        "Market": top5["title"],
        "Ticker": top5["ticker"],
        "Trading Volume": top5["volume_fp"].round(0).astype("int64"),
        "Kalshi Price": top5["yes_ask_dollars"],
    })

    totals = (
        raw.groupby("category")["volume_fp"]
        .agg(Markets_Found="count", Total_Volume="sum")
        .reset_index()
        .rename(columns={"category": "Category"})
    )
    top5_vol = top5.groupby("category")["volume_fp"].sum().reset_index()
    top5_vol.columns = ["Category", "Top5_Volume"]
    totals = totals.merge(top5_vol, on="Category", how="left")
    totals["Total_Volume"] = totals["Total_Volume"].round(0).astype("int64")
    totals["Top5_Volume"] = totals["Top5_Volume"].fillna(0).round(0).astype("int64")

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        export_df.to_excel(writer, sheet_name="Analysis", index=False)
        verify.to_excel(writer, sheet_name="Volume Verification", index=False)
        totals.to_excel(writer, sheet_name="Category Totals", index=False)

        # Real numeric currency formatting so volumes are sortable/verifiable
        for sheet, col, width in (("Volume Verification", "E", 18), ("Category Totals", "C", 18)):
            ws = writer.sheets[sheet]
            for cell in ws[col][1:]:
                cell.number_format = '$#,##0'
            ws.column_dimensions[col].width = width
        ws = writer.sheets["Category Totals"]
        for cell in ws["D"][1:]:
            cell.number_format = '$#,##0'
        ws.column_dimensions["D"].width = 18
        vv = writer.sheets["Volume Verification"]
        vv.column_dimensions["C"].width = 60
        vv.column_dimensions["D"].width = 32


def build_substack_brief(df, sections, report_date, output_path):
    lines = [
        f"# AFG Intelligence Brief — {report_date}",
        "",
        "*Where the crowd is wrong, and why.*",
        "",
        "---",
        "",
        sections.get("Executive Summary", ""),
        "",
    ]

    top = df.iloc[0]
    lines.append(f"## The big one: {top['market']}")
    lines.append("")
    lines.append(
        f"Kalshi has this at **{top['kalshi_price']:.0%}**. We're at **{top['afg_probability']:.0%}** "
        f"— a {abs(top['edge_score']*100):.0f}-point gap, our highest-conviction call today. "
        f"{top['recommendation']}."
    )
    lines.append("")

    for category in CATEGORY_ORDER:
        if category in sections and sections[category]:
            lines.append(f"## {category}")
            lines.append("")
            lines.append(sections[category])
            lines.append("")

    lines.append("---")
    lines.append("")
    lines.append("*As always: this is research, not financial advice. Trade your own book.*")
    lines.append("")
    lines.append("— AFG Research")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def main():
    require_approved_files()
    os.makedirs(REPORTS_DIR, exist_ok=True)

    df = load_predictions()
    sections, contrarian_rows = parse_narrative()
    report_date = date.today().strftime("%B %d, %Y")
    file_date = date.today().isoformat()

    docx_path = f"{REPORTS_DIR}/AFG_Research_Report_{file_date}.docx"
    brief_path = f"{REPORTS_DIR}/AFG_Intelligence_Brief_{file_date}_Substack.md"

    build_docx(df, sections, contrarian_rows, report_date, docx_path)
    build_xlsx(df, XLSX_OUTPUT_PATH)
    build_substack_brief(df, sections, report_date, brief_path)

    print(f"Wrote {docx_path}")
    print(f"Wrote {XLSX_OUTPUT_PATH}")
    print(f"Wrote {brief_path}")


if __name__ == "__main__":
    main()
